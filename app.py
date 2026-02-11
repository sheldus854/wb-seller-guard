import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import datetime
import requests
import google.generativeai as genai

# --- 1. НАСТРОЙКИ ---
st.set_page_config(page_title="SellerGuard AI", page_icon="🛡️", layout="wide")

# --- 2. ФУНКЦИИ БАЗЫ (Supabase) ---
def get_secrets():
    try:
        return st.secrets["supabase"]["url"], st.secrets["supabase"]["key"]
    except:
        return None, None

def send_to_supabase(contact, problem, amount):
    url, key = get_secrets()
    if not url: return False
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Content-Type": "application/json", "Prefer": "return=minimal"}
    data = {"contact": contact, "problem_type": problem, "amount": int(amount)}
    try:
        requests.post(f"{url}/rest/v1/leads", headers=headers, json=data)
        return True
    except:
        return False

def fetch_leads():
    url, key = get_secrets()
    if not url: return []
    headers = {"apikey": key, "Authorization": f"Bearer {key}"}
    try:
        r = requests.get(f"{url}/rest/v1/leads?select=*", headers=headers)
        return pd.DataFrame(r.json()) if r.status_code == 200 else pd.DataFrame()
    except:
        return pd.DataFrame()

# --- 3. МОЗГИ (Gemini) ---
def get_ai_response(user_question):
    try:
        api_key = st.secrets["gemini"]["api_key"]
        genai.configure(api_key=api_key)
    except:
        return "⚠️ Ошибка: Нет ключа Gemini в secrets.toml"

    # Читаем базу знаний
    try:
        with open("knowledge.txt", "r", encoding="utf-8") as f:
            knowledge_base = f.read()
    except:
        knowledge_base = "База знаний временно недоступна."

    # ИСПОЛЬЗУЕМ 1.5 FLASH (Самая стабильная и дешевая по лимитам)
    # Если 1.5 не работает, можно заменить на 'gemini-2.0-flash'
    try:
        # Используем модель, которая ЕСТЬ в твоем списке
        model = genai.GenerativeModel('gemini-2.0-flash-lite')
        
        prompt = f"""
        Ты — SellerGuard, профессиональный юрист по защите прав селлеров Wildberries.
        Твоя цель: защитить деньги селлера.
        
        БАЗА ЗНАНИЙ И ЗАКОНЫ:
        {knowledge_base}
        
        Вопрос клиента: {user_question}
        
        Дай четкий, юридически обоснованный ответ. Обязательно ссылайся на статьи ГК РФ или пункты Оферты, если они есть в базе.
        """

        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"🚨 Ошибка Google API: {e}. (Попробуй обновить ключ)"

# --- 4. ГЕНЕРАТОР ДОКУМЕНТОВ ---
def create_doc(seller, inn, act, money, problem):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Шапка
    doc.add_paragraph(f"В ООО «Вайлдберриз» (ОГРН 1067746062449)")
    doc.add_paragraph(f"От: {seller} (ИНН {inn})")
    doc.add_paragraph(f"Дата: {datetime.date.today()}\n")
    
    # Заголовок
    p = doc.add_paragraph("ДОСУДЕБНАЯ ПРЕТЕНЗИЯ")
    p.alignment = 1 # По центру
    
    # Тело
    doc.add_paragraph(f"\nМною был получен Отчет о реализации/Удержаниях № {act}.")
    doc.add_paragraph(f"В данном отчете было произведено необоснованное удержание на сумму {money} руб.")
    doc.add_paragraph(f"Суть нарушения: {problem}.")
    
    doc.add_paragraph("\nОбоснование:")
    doc.add_paragraph("1. Согласно ст. 1102 ГК РФ, лицо, которое без установленных законом оснований приобрело имущество за счет другого лица, обязано возвратить неосновательное обогащение.")
    doc.add_paragraph("2. Wildberries не предоставил доказательств нарушения (фото/видео/акты), подтверждающих правомерность штрафа.")
    
    doc.add_paragraph("\nПРОШУ:")
    doc.add_paragraph(f"1. Отменить удержание по отчету № {act}.")
    doc.add_paragraph(f"2. Вернуть на баланс денежные средства в размере {money} руб. в течение 10 календарных дней.")
    
    doc.add_paragraph("\nВ случае отказа я буду вынужден обратиться в Арбитражный суд с возложением на вас судебных расходов.")
    
    doc.add_paragraph(f"\n_______________ / {seller}")
    
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# --- 5. ИНТЕРФЕЙС ---
with st.sidebar:
    st.header("🔐 Владелец")
    if st.text_input("Пароль", type="password") == st.secrets["admin"]["password"]:
        st.success("Доступ открыт")
        df = fetch_leads()
        if not df.empty:
            st.write("Заявки на юриста:")
            st.dataframe(df)
            st.metric("Потенциал выручки (15%)", f"{int(df['amount'].sum() * 0.15):,} ₽")
        else:
            st.info("Заявок пока нет")

st.title("🛡️ SellerGuard AI")
st.markdown("#### Твой личный юрист и защита от штрафов WB")

tabs = st.tabs(["💬 AI-Консультант", "📄 Генератор Претензий", "👨‍⚖️ Нанять Юриста"])

# Вкладка 1: ЧАТ
with tabs[0]:
    st.info("🤖 Я изучил базу выигранных дел. Спроси меня про штрафы или блокировки.")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if prompt := st.chat_input("Например: 'ВБ потерял товар, как вернуть деньги?'"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Анализирую судебную практику..."):
                reply = get_ai_response(prompt)
                st.markdown(reply)
        st.session_state.messages.append({"role": "assistant", "content": reply})

# Вкладка 2: ГЕНЕРАТОР (ВЕРНУЛСЯ!)
with tabs[1]:
    st.write("### 📝 Создание документа за 1 минуту")
    st.write("Заполни форму, и я сформирую юридически грамотную претензию.")
    
    c1, c2 = st.columns(2)
    with c1:
        name = st.text_input("Ваше название (ИП/ООО)", "ИП Иванов И.И.")
        inn = st.text_input("Ваш ИНН", "500100200300")
    with c2:
        act = st.text_input("Номер Отчета/Штрафа", "№ 12345678")
        money = st.number_input("Сумма ущерба (руб)", 50000, step=1000)
    
    problem = st.text_area("Описание проблемы (кратко)", "ВБ удержал штраф за габариты, но замеры не предоставил. Товар соответствует карточке.")
    
    if st.button("Сгенерировать Претензию"):
        if name and inn and act:
            file_doc = create_doc(name, inn, act, money, problem)
            st.success("Документ готов!")
            st.download_button(
                label="📥 Скачать Претензию (.docx)",
                data=file_doc,
                file_name=f"Pretenziya_{act}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Пожалуйста, заполните ИНН и номер отчета.")

# Вкладка 3: ФОРМА ЛИДОВ (ВЕРНУЛАСЬ!)
with tabs[2]:
    st.write("### 🤝 Сложный случай? Передадим дело профи.")
    st.write("Наши юристы работают за % от выигранной суммы. Оплата только после победы.")
    
    with st.form("lead_form"):
        contact = st.text_input("Ваш контакт (Telegram/WhatsApp)", placeholder="@username или +7...")
        problem_desc = st.text_area("Суть проблемы", placeholder="Штраф 200к, заблокировали кабинет...")
        amount_lost = st.number_input("Сумма спора", min_value=10000, value=100000)
        
        submitted = st.form_submit_button("🚀 Отправить заявку юристу")
        if submitted:
            if contact and problem_desc:
                # Отправка в Supabase
                if send_to_supabase(contact, problem_desc, amount_lost):
                    st.success("Заявка принята! Юрист свяжется с вами в течение часа.")
                else:
                    st.error("Ошибка соединения. Напишите нам напрямую в поддержку.")
            else:
                st.warning("Заполните контакт, чтобы мы могли связаться.")

